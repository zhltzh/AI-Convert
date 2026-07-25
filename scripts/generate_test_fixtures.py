"""Generate local-only regression fixtures for browser conversion checks."""
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen.canvas import Canvas

OUT = Path(__file__).resolve().parents[1] / "test-fixtures"
OUT.mkdir(exist_ok=True)

doc = Document()
doc.add_heading("AI Convert 文档转换验收样本", 0)
doc.add_paragraph("用于验证 Word 转 Markdown 的标题、列表、链接和表格结构。")
doc.add_heading("重点清单", level=1)
for item in ["保留标题层级", "保留项目列表", "保留基础表格"]:
    doc.add_paragraph(item, style="List Bullet")
table = doc.add_table(rows=1, cols=2)
table.rows[0].cells[0].text = "格式"
table.rows[0].cells[1].text = "预期"
for left, right in [("DOCX", "转为 Markdown"), ("CSV", "转为 Markdown 表格")]:
    cells = table.add_row().cells
    cells[0].text, cells[1].text = left, right
doc.add_paragraph("更多信息：https://example.com", style=None)
doc.save(OUT / "conversion-sample.docx")

(OUT / "conversion-sample.csv").write_text("名称,格式,状态\nWord,DOCX,可转换\nPDF,文本型 PDF,可转换\n", encoding="utf-8")

pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
canvas = Canvas(str(OUT / "conversion-sample.pdf"), pagesize=A4)
canvas.setFont("STSong-Light", 20)
canvas.drawString(72, 780, "AI Convert PDF 验收样本")
canvas.setFont("STSong-Light", 12)
canvas.drawString(72, 742, "这是一份包含可选文字的 PDF，用于测试文本提取。")
canvas.drawString(72, 712, "1. 标题、段落和列表应被识别。")
canvas.drawString(72, 688, "2. 第二行用于验证多行文本。")
canvas.save()
